#!/usr/bin/env python3
"""
Export Plan van Aanpak markdown files to one Word document (.docx).

Usage:
    python scripts/export_pva_to_word.py

Output:
    docs/plan-van-aanpak/Plan-van-Aanpak-MSG3-Maximo-Converter.docx
"""

import re
from pathlib import Path

from docx import Document


def get_project_root() -> Path:
    """Project root (where docs/ lives)."""
    script_dir = Path(__file__).resolve().parent
    return script_dir.parent


def parse_markdown_line(line: str) -> tuple[str | None, int | None, str]:
    """
    Return (kind, level, content) for a line.
    kind: 'heading' | 'bullet' | 'table_sep' | 'table_row' | 'paragraph' | 'empty'
    """
    stripped = line.rstrip()
    if not stripped:
        return ("empty", None, "")

    if stripped.startswith("#"):
        match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if match:
            level = len(match.group(1))
            return ("heading", level, match.group(2).strip())
    if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
        return ("bullet", None, stripped)
    if "|" in stripped and stripped.strip().startswith("|"):
        if re.match(r"^[\s|:-]+$", stripped):
            return ("table_sep", None, stripped)
        return ("table_row", None, stripped)
    return ("paragraph", None, stripped)


def add_paragraph_with_inline_format(doc: Document, text: str, style: str | None = None):
    """Add a paragraph, rendering **bold** as bold runs."""
    p = doc.add_paragraph(style=style)
    if not text:
        return p
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2] + " ")
            run.bold = True
        else:
            p.add_run(part if part else "")
    return p


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Parse consecutive table rows; return (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines):
        kind, _, content = parse_markdown_line(lines[i])
        if kind == "table_sep":
            i += 1
            continue
        if kind != "table_row":
            break
        cells = [c.strip() for c in content.split("|") if c.strip() or content.startswith("|")]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def md_to_docx_content(doc: Document, md_text: str):
    """Append markdown content to document."""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        kind, level, content = parse_markdown_line(line)

        if kind == "empty":
            i += 1
            continue
        if kind == "heading":
            if level == 1:
                doc.add_heading(content, level=0)
            else:
                doc.add_heading(content, level=min(level, 3))
            i += 1
            continue
        if kind == "bullet":
            text = re.sub(r"^[-*]\s+", "", content)
            text = re.sub(r"^\d+\.\s+", "", text)
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue
        if kind == "table_row":
            table_rows, next_i = parse_table_rows(lines, i)
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                for ri, row_cells in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell_text
                i = next_i
            else:
                i += 1
            continue
        if kind == "paragraph":
            add_paragraph_with_inline_format(doc, content)
            i += 1
            continue
        i += 1


def main():
    root = get_project_root()
    docs_dir = root / "docs" / "plan-van-aanpak"
    out_path = docs_dir / "Plan-van-Aanpak-MSG3-Maximo-Converter.docx"

    doc = Document()
    doc.add_heading("Plan van Aanpak – MSG-3 to Maximo Converter", 0)
    doc.add_paragraph(
        "Babcock Schiphol | Pedro Eduardo Cardoso | Associate Degree Software Developer (ADSD), Windesheim"
    )
    doc.add_paragraph(
        "Dit document bevat het volledige plan van aanpak: projectaanpak, planning, risicoanalyse, randvoorwaarden en deliverables."
    )
    doc.add_paragraph()

    order = [
        "01-projectaanpak.md",
        "02-planning.md",
        "03-risicoanalyse.md",
        "04-randvoorwaarden.md",
        "05-deliverables.md",
    ]

    for filename in order:
        path = docs_dir / filename
        if not path.exists():
            print(f"Warning: {path} not found, skipping.")
            continue
        md_text = path.read_text(encoding="utf-8")
        md_to_docx_content(doc, md_text)
        doc.add_paragraph()
        if filename != order[-1]:
            doc.add_page_break()

    doc.save(str(out_path))
    print(f"Word document saved: {out_path}")


if __name__ == "__main__":
    main()
