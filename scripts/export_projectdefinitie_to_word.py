#!/usr/bin/env python3
"""
Export projectdefinitie markdown files to one Word document (.docx).

De volgorde van secties komt uit docs/projectdefinitie/word-export-order.txt
(template-gestuurd; zie PROJECTDEFINITIE-STRUCTUUR.md).

Usage:
    python scripts/export_projectdefinitie_to_word.py

Output:
    docs/projectdefinitie/Projectdefinitie-MSG3-Maximo-Converter.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def get_project_root() -> Path:
    """Project root (where docs/ lives)."""
    script_dir = Path(__file__).resolve().parent
    return script_dir.parent


def parse_markdown_line(line: str) -> tuple[str | None, int | None, str]:
    """
    Return (kind, level, content) for a line.
    kind: 'heading' | 'bullet' | 'table_sep' | 'table_row' | 'paragraph' | 'empty'
    level: for heading 1-6, else None
    """
    stripped = line.rstrip()
    if not stripped:
        return ("empty", None, "")

    # Headings (Markdown supports 1-6 levels; Word typically uses 0-3)
    if stripped.startswith("#"):
        match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if match:
            level = len(match.group(1))
            return ("heading", level, match.group(2).strip())
    # Unordered list
    if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
        return ("bullet", None, stripped)
    # Table row (contains |)
    if "|" in stripped and stripped.strip().startswith("|"):
        if re.match(r"^[\s|:-]+$", stripped):
            return ("table_sep", None, stripped)
        return ("table_row", None, stripped)
    return ("paragraph", None, stripped)


def add_paragraph_with_inline_format(doc: Document, text: str, style: str | None = None):
    """Add a paragraph, rendering **bold** as bold runs."""
    p = doc.add_paragraph(style=style)
    if not text:
        return p
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part if part else "")
    return p


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Parse consecutive table rows; return (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines):
        kind, _, content = parse_markdown_line(lines[i])
        if kind == "table_sep":
            i += 1
            continue
        if kind != "table_row":
            break
        cells = [c.strip() for c in content.split("|") if c.strip() or content.startswith("|")]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def md_to_docx_content(doc: Document, md_text: str):
    """Append markdown content to document."""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        kind, level, content = parse_markdown_line(line)

        if kind == "empty":
            i += 1
            continue
        if kind == "heading":
            if level == 1:
                doc.add_heading(content, level=0)
            else:
                doc.add_heading(content, level=min(level, 3))
            i += 1
            continue
        if kind == "bullet":
            text = re.sub(r"^[-*]\s+", "", content)
            text = re.sub(r"^\d+\.\s+", "", text)
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue
        if kind == "table_row":
            table_rows, next_i = parse_table_rows(lines, i)
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                for ri, row_cells in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell_text
                i = next_i
            else:
                i += 1
            continue
        if kind == "paragraph":
            add_paragraph_with_inline_format(doc, content)
            i += 1
            continue
        i += 1


def load_export_order(docs_dir: Path, default_order: list[str]) -> list[str]:
    """
    Lees exportvolgorde uit word-export-order.txt in docs_dir (template-gestuurd).
    Regels die met # beginnen of leeg zijn worden overgeslagen.
    Als het bestand ontbreekt of leeg is, wordt default_order gebruikt.
    """
    manifest = docs_dir / "word-export-order.txt"
    if not manifest.exists():
        return default_order
    lines = [
        line.strip()
        for line in manifest.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    return lines if lines else default_order


def main():
    root = get_project_root()
    docs_dir = root / "docs" / "projectdefinitie"
    out_path = docs_dir / "Projectdefinitie-MSG3-Maximo-Converter.docx"

    # Geconsolideerd: 1 bestand per deliverable
    src_path = docs_dir / "projectdefinitie.md"
    if src_path.exists():
        doc = Document()
        doc.add_heading("Projectdefinitie – MSG-3 to Maximo Converter", 0)
        doc.add_paragraph(
            "Babcock Schiphol | Pedro Eduardo Cardoso | Associate Degree Software Developer (ADSD), Windesheim"
        )
        doc.add_paragraph()
        md_text = src_path.read_text(encoding="utf-8")
        md_to_docx_content(doc, md_text)
    else:
        # Fallback: oude structuur met meerdere bestanden
        default_order = [
            "00-inleiding.md",
            "01-projectdefinitie.md",
            "02-het-project.md",
            "03-projectaanpak.md",
            "04-projectmanagementorganisatie.md",
            "05-managementstrategieen.md",
            "06-planning.md",
            "07-literatuurlijst.md",
            "08-bijlagen.md",
        ]
        order = load_export_order(docs_dir, default_order)
        doc = Document()
        doc.add_heading("Projectdefinitie – MSG-3 to Maximo Converter", 0)
        doc.add_paragraph(
            "Babcock Schiphol | Pedro Eduardo Cardoso | Associate Degree Software Developer (ADSD), Windesheim"
        )
        doc.add_paragraph()
        for filename in order:
            path = docs_dir / filename
            if not path.exists():
                print(f"Warning: {path} not found, skipping.")
                continue
            md_text = path.read_text(encoding="utf-8")
            md_to_docx_content(doc, md_text)
            doc.add_paragraph()
            if filename != order[-1]:
                doc.add_page_break()

    doc.save(str(out_path))
    print(f"Word document saved: {out_path}")


if __name__ == "__main__":
    main()
